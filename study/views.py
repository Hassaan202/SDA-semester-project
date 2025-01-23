from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from .models import *
from django.urls import reverse
from django.core.files.storage import default_storage
import os 
from django.conf import settings
import pandas as pd

# Create your views here.
def generate_question_paper(file_path):
    # Load the refined dataset
    data = pd.read_csv(file_path)

    # Generate the paper with randomized selection
    paper_questions = []
    chapters = data['Chapter_Number'].unique()

    for chapter in chapters:
        chapter_questions = data[data['Chapter_Number'] == chapter]

        # Randomly select one question from the chapter
        selected_row = chapter_questions.sample(1).iloc[0]  # Randomly select one row
        question_text = selected_row['Improved_Question_Text']
        original_text = selected_row['Question_Text']

        # Ensure no invalid questions in the final paper
        if "True" in question_text or "False" in question_text or not question_text.strip():
            question_text = original_text  # Fallback to the original question if invalid

        # Split the question into words
        words = question_text.split()

        # Remove the first word
        words = words[1:]

        # Check if the next word starts with 'P', and remove it if true
        if len(words) > 0 and words[0][0].upper() == 'P':
            words = words[1:]

        # Join the remaining words back into the question
        question_text = ' '.join(words)

        # Add the question to the paper
        paper_questions.append(f"Chapter {chapter}: {question_text}")

    # Join all questions into a single content string
    return "\n\n".join(paper_questions)


# @login_required
# def aiPage(request):
#     # Path to the CSV file
#     file_path = os.path.join(settings.MEDIA_ROOT, 'refined_chapter_questions.csv')

#     # Check if file exists
#     if not os.path.exists(file_path):
#         return render(request, "aiPage.html", {"error": "Dataset not found. Please ensure the CSV file is uploaded."})

#     # Generate the question paper
#     paper_content = generate_question_paper(file_path).strip()

#     return render(request, "aiPage.html", {"paper_content": paper_content})



@login_required
def resourceSharing(request):
    subjects = Subject.objects.all()  # Retrieve all subjects from the database

    if request.method == 'POST':
        subject_id = request.POST.get('subject')  # Get subject ID from form data
        uploaded_file = request.FILES['file']  # Get uploaded file

        # Get the subject instance
        subject = Subject.objects.get(id=subject_id)

        # Create a new ResourceUpload instance without saving yet
        resource_upload = ResourceUpload(user=request.user, subject=subject)

        # Save file to the specific subject folder using model's save method
        resource_upload.file.save(uploaded_file.name, uploaded_file)

        messages.success(request, "File uploaded successfully!")  # Display success message

    # Get all resources for each subject to display after upload
    resources_by_subject = {}
    for subject in subjects:
        resources_by_subject[subject] = ResourceUpload.objects.filter(subject=subject)

    return render(request, 'resourceSharing.html', {
        'subjects': subjects,
        'resources_by_subject': resources_by_subject,
    })



@login_required
def resourcesLib(request):
    subList=Subject.objects.all()
    selected_subject = None
    playLists = []
    books = []
    pastPapers = []
    
    if request.method == 'POST':
        subject_id = request.POST.get('subject')
        selected_subject = Subject.objects.get(id=subject_id) if subject_id else None

        if selected_subject:
            playLists = ResourcesLib.objects.filter(subject=selected_subject, fileType='playlists')
            books = ResourcesLib.objects.filter(subject=selected_subject, fileType='books')
            pastPapers = ResourcesLib.objects.filter(subject=selected_subject, fileType='pastpapers')

    return render(request, "subjectResources.html", {
        "subList": subList,
        "selected_subject": selected_subject,
        "playLists": playLists,
        "books": books,
        "pastPapers": pastPapers,
    })



@login_required
def discussion_forum(request):
    subjects = Subject.objects.all()
    selected_subject = None
    discussions = []
    
    if request.method == 'POST':
        subject_id = request.POST.get('subject')
        selected_subject = Subject.objects.get(id=subject_id) if subject_id else None
        
        if 'discussion_content' in request.POST and selected_subject:
            # Create a new discussion with the current logged-in user
            discussion_content = request.POST.get('discussion_content')
            new_discussion = Discussion(subj_name=selected_subject, content=discussion_content, user_ID=request.user)

            messages.success(request, "Your query has been succesfully posted!")
            new_discussion.save()
        
        if selected_subject:
            discussions = Discussion.objects.filter(subj_name=selected_subject)
    
    return render(request, 'forum.html', {
        'subjects': subjects,
        'selected_subject': selected_subject,
        'discussions': discussions,
    })







#INTEGRATING AI IN THE APPLICATION

import google.generativeai as genai
import re
import docx
import pdfplumber
from tabulate import tabulate

# Step 1: Find the subject folder based on user input
def find_subject_folder(subject_name, base_dir='past_papers'):
    for folder in os.listdir(base_dir):
        if folder.lower() == subject_name.lower():
            return os.path.join(base_dir, folder)
    return None


# Step 2: Fetch all PDF and DOCX files from the selected folder
def get_files_from_folder(folder_path):
    files = []
    if folder_path and os.path.exists(folder_path):
        for file in os.listdir(folder_path):
            if file.endswith(".pdf") or file.endswith(".docx"):
                files.append(os.path.join(folder_path, file))
    return files


# Step 3: Extract text and tables from PDF files
def extract_text_and_tables_from_pdf(file_path):
    all_content = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Extracting text from the page
            text = page.extract_text()
            if text:
                all_content.append(('text', clean_extracted_text(text)))
            # Extracting tables from the page
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    # Convert table to a string to check for metadata keywords
                    # Ensure to filter out None values
                    table_str = " ".join([" ".join(filter(None, row)) for row in table])
                    # Check for common metadata keywords
                    if not re.search(
                            r'(Course Name|Course Code|Semester|Program|Duration|Total Marks|Paper Date|Exam Type|Section|Page\(s\))',
                            table_str, re.IGNORECASE):
                        formatted_table = tabulate(table, tablefmt="grid")
                        all_content.append(('table', formatted_table))
    return all_content


# Step 4: Extract text and tables from DOCX files
def extract_text_and_tables_from_docx(file_path):
    all_content = []
    doc = docx.Document(file_path)

    # Extracting text paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_content.append(('text', clean_extracted_text(text)))

    # Extracting tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        # Convert table data to a single string to check for metadata keywords
        table_str = " ".join([" ".join(filter(None, row)) for row in table_data])
        if not re.search(
                r'(Course Name|Course Code|Semester|Program|Duration|Total Marks|Paper Date|Exam Type|Section|Page\(s\))',
                table_str, re.IGNORECASE):
            formatted_table = tabulate(table_data, tablefmt="grid")
            all_content.append(('table', formatted_table))

    return all_content


# Step 5: Clean extracted text to remove headers, footers, and irrelevant data
def clean_extracted_text(text):
    # Remove unnecessary details like page numbers, headers, and metadata
    text = re.sub(
        r'Page \d+|Roll No\.|Name|Section|Semester|Duration|Total Marks|Paper Date|Course|Exam|Weight|Scratch sheet|rough work.*|Department.*|National University.*',
        '', text, flags=re.IGNORECASE)
    text = re.sub(r'\d{1,2}/\d{1,2}/\d{2,4}', '', text)  # Remove dates
    text = re.sub(r'\s+', ' ', text)  # Normalize whitespaces
    return text.strip()


# Step 6: Process all files in the folder to extract text and tables together
def process_all_files(file_paths):
    all_content = []
    for idx, file_path in enumerate(file_paths, 1):
        content = []
        if file_path.endswith(".pdf"):
            content = extract_text_and_tables_from_pdf(file_path)
        elif file_path.endswith(".docx"):
            content = extract_text_and_tables_from_docx(file_path)

        all_content.append((f"Paper No {idx}:", content))  # Store paper number with content

    return all_content


# Step 7: Extract all questions and associate them with tables
def extract_questions_and_tables(papers_content):
    questions_with_tables = []

    for paper_title, content in papers_content:
        current_question = None

        for item_type, item_content in content:
            if item_type == 'text':
                # Extract questions from text using a regex that matches questions starting with "Q" followed by digits
                potential_questions = re.findall(r'(Q\d+\.\s*.*?)(?=(?:Q\d+\.|\Z))', item_content, re.DOTALL)

                if potential_questions:
                    for question in potential_questions:
                        if current_question:
                            questions_with_tables.append(
                                (paper_title, current_question.strip()))  # Save the previous question with tables
                        current_question = clean_question(question)  # Start a new question
                else:
                    if current_question:
                        current_question += "\n" + item_content  # Append any additional text to the current question

            elif item_type == 'table' and current_question:
                # Attach the table to the current question if there is one
                current_question += "\n" + item_content

        if current_question:
            questions_with_tables.append(
                (paper_title, current_question.strip()))  # Append the last question of each paper

    return questions_with_tables


# Step 8: Clean individual questions to remove unnecessary content
def clean_question(question):
    question = re.sub(r'(?:Instructions|NOTE:).*', '', question, flags=re.IGNORECASE)
    return question.strip()


# Step 9: Generate a list of all questions with tables
def generate_question_list(questions_with_tables):
    question_list = []
    for paper_title, question in questions_with_tables:
        question_list.append(f"{paper_title}\n{question.strip()}")
    return "\n".join(question_list)


def getAIresponse(prompt):
    genai.configure(api_key="HIDDEN")
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt)
    return response.text


@login_required
def aiPage(request):
    paper_content = ""
    generated_questions = ""
    error_message = ""
    subList = Subject.objects.all()
    selected_subject = None

    if request.method == "POST":
        subject_id = request.POST.get("subject")
        selected_subject = Subject.objects.get(id=subject_id) if subject_id else None
        subject_name = selected_subject.name if selected_subject else None

        # Check and find the subject folder
        subject_folder = find_subject_folder(subject_name)

        # Initialize variables for AI generation
        generated_questions = None
        error_message = None

        # Step 1: Check if the subject folder exists
        if subject_folder:
            subject_files = get_files_from_folder(subject_folder)

            # Step 2: Process files if available
            if subject_files:
                papers_content = process_all_files(subject_files)

                # Step 3: Extract questions and tables
                questions_with_tables = extract_questions_and_tables(papers_content)
                question_list = generate_question_list(questions_with_tables)

                # Step 4: Generate AI-based questions using the extracted questions
                max_questions = 5
                try:
                    # Generate similar questions using AI
                    generated_questions = getAIresponse(
                        f"Generate {max_questions} questions based on the following topics. "
                        f"Provide only the questions, nothing else:\n{question_list}"
                    )
                except Exception as e:
                    error_message = f"Error generating AI questions: {str(e)}"
            else:
                error_message = f"No files found in the subject folder '{subject_name}'."
        else:
            error_message = f"Subject folder '{subject_name}' not found."

        # Step 5: Generate the question paper from the CSV dataset
        file_path = os.path.join(settings.MEDIA_ROOT, "refined_chapter_questions.csv")
        if os.path.exists(file_path):
            paper_content = generate_question_paper(file_path).strip()
        else:
            error_message = error_message or "Dataset not found. Please ensure the CSV file is uploaded."

    # Step 6: Pass data to the template for rendering
    return render(
        request,
        "aiPage.html",
        {
            "paper_content": paper_content,
            "generated_questions": generated_questions,
            "error": error_message,
            "subList": subList,
            "selected_subject": selected_subject,
        },
    )



